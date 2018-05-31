from docx import Document
import os
files = os.listdir()
res = Document()
for file in files:
    if os.path.splitext(file)[1] == '.docx':
        d = Document(open(file, 'rb'))
        for par in d.paragraphs:
            res.add_paragraph(par.text)
res.save('test.docx')
